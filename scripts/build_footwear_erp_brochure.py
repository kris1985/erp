from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

OUT = "docs/鞋厂定制ERP-AI企微三折页.docx"

NAVY = "102A43"
BLUE = "176B87"
TEAL = "1D8B8A"
GOLD = "E7B74B"
LIGHT = "EFF6F7"
PALE = "F7FAFA"
MUTED = "5E6B76"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="FFFFFF", size="0"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=180, start=220, bottom=180, end=220):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=10, color=NAVY, bold=False, font="Hiragino Sans GB"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def paragraph(cell, text="", size=10, color=NAVY, bold=False, after=5, before=0, align=None, line=1.15):
    p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size, color, bold)
    return p


def title(cell, text, cover=False):
    p = paragraph(cell, text, 25 if cover else 16, WHITE if cover else NAVY, True, after=8, line=1.0)
    return p


def kicker(cell, text, cover=False):
    return paragraph(cell, text, 8.5, GOLD if cover else TEAL, True, after=8)


def bullets(cell, items, size=9.4, color=NAVY):
    for item in items:
        p = cell.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        set_run(p.add_run(item), size, color)


def band(cell, label, detail, dark=False):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    a = p.add_run(label + "  ")
    set_run(a, 9, GOLD if dark else TEAL, True)
    b = p.add_run(detail)
    set_run(b, 9, WHITE if dark else NAVY, False)


def clear_cell(cell):
    p = cell.paragraphs[0]
    p.clear()


def build_page(doc, fills):
    table = doc.add_table(rows=1, cols=3)
    set_table_fixed(table, [3250, 3250, 3250])
    for cell, fill in zip(table.rows[0].cells, fills):
        clear_cell(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_shading(cell, fill)
        set_cell_border(cell)
        set_cell_margins(cell)
    return table.rows[0].cells


def add_footer(cell, page_note, dark=False):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(page_note), 7.5, "C9D7DA" if dark else MUTED)


# The bundled headless LibreOffice runtime does not ship a CJK font.  For this
# print-first brochure we draw the two full-bleed A4 panels with the host's
# PingFang font, then place the high-resolution artwork into the DOCX.
CANVAS = (3508, 2480)  # A4 landscape at 300 dpi
PANEL_W = CANVAS[0] // 3
FONT_PATH = "/System/Library/Fonts/Hiragino Sans GB.ttc"


def gfont(size, bold=False):
    return ImageFont.truetype(FONT_PATH, size=size, index=1 if bold else 0)


def wrapped(draw, text, font, width):
    lines, line = [], ""
    for char in text:
        candidate = line + char
        if line and draw.textlength(candidate, font=font) > width:
            lines.append(line)
            line = char
        else:
            line = candidate
    if line:
        lines.append(line)
    return lines


def multiline(draw, xy, text, font, fill, width, leading=1.3, gap=0):
    x, y = xy
    lines = []
    for raw in text.split("\n"):
        lines.extend(wrapped(draw, raw, font, width) or [""])
    step = int(font.size * leading)
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += step
    return y + gap


def brochure_page(specs, path):
    image = Image.new("RGB", CANVAS, "#FFFFFF")
    draw = ImageDraw.Draw(image)
    for idx, spec in enumerate(specs):
        x0 = idx * PANEL_W
        bg = spec["bg"]
        ink = spec.get("ink", NAVY)
        accent = spec.get("accent", TEAL)
        draw.rectangle((x0, 0, x0 + PANEL_W, CANVAS[1]), fill=f"#{bg}")
        if idx in (1, 2):
            draw.line((x0, 0, x0, CANVAS[1]), fill="#D9E6E8", width=3)
        x, y, width = x0 + 92, 122, PANEL_W - 184
        draw.text((x, y), spec["kicker"], font=gfont(30, True), fill=f"#{accent}")
        y += 88
        for line in spec["title"].split("\n"):
            draw.text((x, y), line, font=gfont(spec.get("title_size", 74), True), fill=f"#{ink}")
            y += int(spec.get("title_size", 74) * 1.16)
        y += 28
        y = multiline(draw, (x, y), spec["body"], gfont(32), f"#{ink}", width, 1.45, 42)
        for label, detail in spec.get("bands", []):
            box_h = 112 if len(detail) < 36 else 142
            fill = "#173B52" if bg == NAVY else "#FFFFFF"
            outline = "#315A71" if bg == NAVY else "#CFE0E4"
            draw.rounded_rectangle((x, y, x + width, y + box_h), radius=18, fill=fill, outline=outline, width=2)
            draw.text((x + 24, y + 18), label, font=gfont(28, True), fill=f"#{GOLD if bg == NAVY else accent}")
            multiline(draw, (x + 24, y + 55), detail, gfont(25), "#FFFFFF" if bg == NAVY else f"#{ink}", width - 48, 1.2)
            y += box_h + 20
        for bullet in spec.get("bullets", []):
            draw.ellipse((x, y + 14, x + 12, y + 26), fill=f"#{accent}")
            y = multiline(draw, (x + 30, y), bullet, gfont(27), f"#{ink}", width - 30, 1.28, 14)
        note = spec.get("note")
        if note:
            y += 12
            multiline(draw, (x, y), note, gfont(25, True), f"#{accent}", width, 1.3)
        footer = spec.get("footer")
        if footer:
            draw.text((x, CANVAS[1] - 102), footer, font=gfont(21), fill="#B9CDD1" if bg == NAVY else f"#{MUTED}")
    image.save(path, quality=95)


def build_artwork():
    outside = [
        dict(bg=PALE, kicker="鞋厂生产执行 ERP", title="让每一张单\n都能跑到交付", body="从接单、齐套、开裁、报工计件，到排产、装箱、出货与追溯，把鞋厂每天最容易失控的环节接成可确认的链路。", bands=[("适用对象", "鞋类 OEM / 外贸工厂 / 多订单快反生产"), ("部署方式", "按工厂业务流程配置；企微机器人接入可按权限范围评估")], note="预约产品演示\n请联系您的产品顾问，获取行业方案与现场演示。", footer="外侧 · 背封"),
        dict(bg=LIGHT, kicker="一套账本 · 一组提醒", title="把问题推到\n该看的人手上", body="不必天天守着系统：缺料、交期风险与进度日报可按租户配置推进企微群机器人。", bands=[("缺料预警", "按订单、物料、预计齐套日组织摘要"), ("交期风险", "提前暴露受影响订单与风险等级"), ("进度日报", "产量 KPI + 重点订单，作为每日生产心跳")], note="预警只推不改；发送失败不阻断业务；群里与系统使用同一套数据口径。", footer="外侧 · 折页"),
        dict(bg=NAVY, ink=WHITE, accent=GOLD, kicker="鞋厂定制 ERP + AI Agent + 企业微信", title="少靠吼，\n少对 Excel", body="鞋厂每天吵的事，系统先算清；该预警的，推到企微；需要拍板的，人再确认。", bands=[("鞋厂定制", "按码算料 · 齐套争料 · 流转卡 · 报工计件 · 排产 · 齐码出货"), ("AI Agent", "把风险、数据与建议组织成可追溯的行动"), ("企微联动", "预警 / 日报通知 · MCP 机器人只读问数")], note="AI 给建议，账本听人的。", footer="外侧 · 封面", title_size=80),
    ]
    inside = [
        dict(bg=WHITE, kicker="01 · 先像鞋厂", title="一条从接单\n到交付的执行链", body="不是通用进销存拼出来的生产页，而是围绕鞋厂色码、流转、计件与齐码交付的日常。", bullets=["接单先看毛利粗估、缺料、交期冲击与争料", "按码算料，库存分配互斥，避免两单同时“假齐套”", "流转卡扫码报工，计件单价锁定，工资可核对", "正排 / 倒排出方案；插单先看影响，再由 PMC 确认", "齐码出货、装箱箱唛、来料 IQC、返修与追溯"], note="主业务不依赖 AI；即使关闭 AI，订单、生产与交付闭环仍照常运行。", footer="内侧 · 鞋厂业务闭环"),
        dict(bg=PALE, kicker="02 · AI 是参谋，不是第二套账", title="把“该问谁”\n变成“先做什么”", body="AI Agent 在规则、权限、证据与审批边界内工作：理解问题、编排只读分析、解释风险、生成建议或草稿。", bands=[("今日行动", "把缺料、交期、进度等风险收敛成优先处理事项"), ("自然语言问数", "用业务语言查回款、利润、订单进度、产能与风险"), ("质量 / 损耗预警", "从工序异常、返修和用料差异中给出抽检或核对建议"), ("排产方案辅助", "比较插单、加班、外协等方案影响，人工确认后落库")], note="所有数值可追溯至指标与证据；AI 不直接改数量、交期、排产、账款或工资。", footer="内侧 · AI Agent 能力", title_size=62),
        dict(bg=LIGHT, kicker="03 · 企业微信成为生产入口", title="群里收预警，\n对话拿结论", body="让老板、PMC、采购与车间主管在企业微信里更快得到同一套事实，而不是在群里反复催问、手工导表。", bands=[("预警 / 日报通知", "企微群机器人接收缺料、交期风险与进度日报；消息可回到 ERP 查看上下文"), ("AI 对话接入企微", "企微智能机器人可经 MCP 调用 ERP 的受限只读工具，以自然语言查询授权数据"), ("可控的安全边界", "MCP 密钥、权限范围与租户隔离；白名单指标；不开放任意 SQL、不让机器人直接写账")], note="企业微信连接的是受控服务，而不是无边界聊天机器人。", footer="内侧 · 企微通知与 AI 对话", title_size=66),
    ]
    brochure_page(outside, "/private/tmp/footwear_erp_outside.png")
    brochure_page(inside, "/private/tmp/footwear_erp_inside.png")


def main():
    build_artwork()
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture("/private/tmp/footwear_erp_outside.png", width=section.page_width, height=section.page_height)
    p.add_run().add_break(WD_BREAK.PAGE)
    p.add_run().add_picture("/private/tmp/footwear_erp_inside.png", width=section.page_width, height=section.page_height)

    doc.save(OUT)


if __name__ == "__main__":
    main()
